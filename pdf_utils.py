import io
import os
import tempfile
import PyPDF2
import fitz  # PyMuPDF
from pdfminer.high_level import extract_text
import docx2pdf
import docx
import textract
import re
from PIL import Image
import pytesseract
import numpy as np
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

def extract_text_from_pdf(file):
    """
    Extract text content from a PDF file with enhanced extraction capabilities
    
    Args:
        file: File buffer from Streamlit file uploader
    
    Returns:
        str: Extracted text from the PDF
        dict: Additional metadata (if available)
    """
    # Reset file pointer
    file.seek(0)
    
    # Get the filename and extension
    filename = file.name if hasattr(file, 'name') else 'uploaded_file'
    extension = os.path.splitext(filename)[1].lower()
    
    # Initialize metadata
    metadata = {
        "pages": 0,
        "file_type": extension,
        "original_filename": filename
    }
    
    # Convert non-PDF files to PDF if needed
    if extension != '.pdf':
        return convert_and_extract(file, extension, metadata)
    
    # Try different extraction methods for PDFs
    # Method 1: PyPDF2
    try:
        text = ""
        pdf_reader = PyPDF2.PdfReader(file)
        metadata["pages"] = len(pdf_reader.pages)
        
        # Extract document info
        if pdf_reader.metadata:
            for key, value in pdf_reader.metadata.items():
                if key.startswith('/'):
                    clean_key = key[1:].lower()
                    metadata[clean_key] = value
        
        # Extract text from each page
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        # If we got substantial text, return it
        if len(text.strip()) > 200:
            return text, metadata
    except Exception as e:
        print(f"PyPDF2 extraction failed: {str(e)}")
    
    # Method 2: PyMuPDF (fitz)
    try:
        file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file.getvalue())
            temp_path = temp_file.name
        
        doc = fitz.open(temp_path)
        metadata["pages"] = doc.page_count
        
        text = ""
        for page in doc:
            text += page.get_text() + "\n\n"
        
        # Clean up the temp file
        doc.close()
        os.unlink(temp_path)
        
        # If we got substantial text, return it
        if len(text.strip()) > 200:
            return text, metadata
    except Exception as e:
        print(f"PyMuPDF extraction failed: {str(e)}")
        # Clean up if file exists
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.unlink(temp_path)
    
    # Method 3: PDFMiner
    try:
        file.seek(0)
        text = extract_text(io.BytesIO(file.getvalue()))
        
        # If we got substantial text, return it
        if len(text.strip()) > 200:
            return text, metadata
    except Exception as e:
        print(f"PDFMiner extraction failed: {str(e)}")
    
    # Method 4: Textract as a last resort
    try:
        file.seek(0)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            temp_file.write(file.getvalue())
            temp_path = temp_file.name
        
        text = textract.process(temp_path, method='pdfminer').decode('utf-8')
        
        # Clean up
        os.unlink(temp_path)
        
        return text, metadata
    except Exception as e:
        print(f"Textract extraction failed: {str(e)}")
        # Clean up if file exists
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.unlink(temp_path)
    
    # If all methods failed, try OCR
    try:
        file.seek(0)
        text = extract_with_ocr(file)
        if text:
            metadata["extraction_method"] = "ocr"
            return text, metadata
    except Exception as e:
        print(f"OCR extraction failed: {str(e)}")
    
    # If all extractions fail
    return "Error: Could not extract text from the document. Please ensure it contains selectable text or try another file.", metadata

def convert_and_extract(file, extension, metadata):
    """
    Convert non-PDF file to PDF and extract text
    
    Args:
        file: File buffer
        extension: File extension
        metadata: Metadata dictionary to update
    
    Returns:
        str: Extracted text
        dict: Updated metadata
    """
    file.seek(0)
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as temp_input:
        temp_input.write(file.getvalue())
        input_path = temp_input.name
    
    output_path = input_path.replace(extension, '.pdf')
    metadata["converted_from"] = extension
    
    try:
        # Handle different file types
        if extension in ['.docx', '.doc']:
            # Convert Word to PDF
            if extension == '.docx':
                doc = docx.Document(input_path)
                docx2pdf.convert(input_path, output_path)
            else:
                # For .doc, use textract to extract directly
                text = textract.process(input_path).decode('utf-8')
                metadata["extraction_method"] = "textract"
                os.unlink(input_path)
                return text, metadata
                
        elif extension in ['.txt', '.rtf', '.odt']:
            # For text files, extract directly
            text = textract.process(input_path).decode('utf-8')
            metadata["extraction_method"] = "textract"
            os.unlink(input_path)
            return text, metadata
            
        elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            # For images, use OCR
            text = extract_with_ocr(file)
            metadata["extraction_method"] = "ocr"
            os.unlink(input_path)
            return text, metadata
        
        # Extract text from the converted PDF
        with open(output_path, 'rb') as pdf_file:
            pdf_content = pdf_file.read()
            pdf_io = io.BytesIO(pdf_content)
            text, _ = extract_text_from_pdf(pdf_io)
        
        # Clean up temporary files
        os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)
        
        return text, metadata
    
    except Exception as e:
        # Clean up on failure
        if os.path.exists(input_path):
            os.unlink(input_path)
        if os.path.exists(output_path):
            os.unlink(output_path)
        
        return f"Error converting {extension} to PDF: {str(e)}", metadata

def extract_with_ocr(file):
    """
    Extract text from an image or scanned PDF using OCR
    
    Args:
        file: File buffer
    
    Returns:
        str: Extracted text
    """
    file.seek(0)
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
        temp_file.write(file.getvalue())
        temp_path = temp_file.name
    
    try:
        # Open PDF with PyMuPDF
        doc = fitz.open(temp_path)
        text = ""
        
        # Process each page
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            
            # Convert page to image
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Apply image preprocessing to improve OCR
            img = preprocess_image(img)
            
            # Extract text with OCR
            page_text = pytesseract.image_to_string(img)
            text += page_text + "\n\n"
        
        # Clean up
        doc.close()
        os.unlink(temp_path)
        
        return text
    
    except Exception as e:
        # Clean up on failure
        if 'doc' in locals():
            doc.close()
        if os.path.exists(temp_path):
            os.unlink(temp_path)
        
        raise e

def preprocess_image(img):
    """
    Preprocess image to improve OCR quality
    
    Args:
        img: PIL Image
    
    Returns:
        PIL Image: Processed image
    """
    # Convert to numpy array
    img_np = np.array(img)
    
    # Convert to grayscale if it's not already
    if len(img_np.shape) == 3:
        img_gray = np.dot(img_np[...,:3], [0.2989, 0.5870, 0.1140])
    else:
        img_gray = img_np
    
    # Apply thresholding to make text stand out
    img_binary = (img_gray > 128).astype(np.uint8) * 255
    
    # Convert back to PIL Image
    return Image.fromarray(img_binary.astype(np.uint8))

def analyze_resume_structure(text):
    """
    Analyze the structure of a resume to identify sections
    
    Args:
        text: Extracted text from resume
    
    Returns:
        dict: Identified sections and their positions
    """
    # Common section headers in resumes
    section_patterns = {
        "personal_info": r"(?i)(personal\s+information|contact|profile)",
        "summary": r"(?i)(summary|objective|professional\s+summary|about\s+me)",
        "education": r"(?i)(education|academic|qualifications|degrees)",
        "experience": r"(?i)(experience|work\s+experience|employment|work\s+history)",
        "skills": r"(?i)(skills|technical\s+skills|competencies|expertise)",
        "projects": r"(?i)(projects|personal\s+projects|academic\s+projects)",
        "certifications": r"(?i)(certifications|certificates|credentials)",
        "awards": r"(?i)(awards|honors|achievements)",
        "languages": r"(?i)(languages|language\s+proficiency)",
        "publications": r"(?i)(publications|papers|research)",
        "references": r"(?i)(references|referees)"
    }
    
    sections = {}
    lines = text.split('\n')
    
    current_section = None
    section_content = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        
        # Check if this line is a section header
        for section, pattern in section_patterns.items():
            if re.search(pattern, line) and (len(line) < 50 or re.match(r'^[A-Z\s]+$', line)):
                # If we were in a section, save it before starting a new one
                if current_section:
                    sections[current_section] = {
                        "content": '\n'.join(section_content),
                        "line_start": sections[current_section].get("line_start"),
                        "line_end": i - 1
                    }
                
                # Start new section
                current_section = section
                section_content = []
                sections[current_section] = {"line_start": i}
                break
        
        # If we're in a section, add this line to the content
        if current_section:
            section_content.append(line)
    
    # Save the last section
    if current_section and section_content:
        sections[current_section] = {
            "content": '\n'.join(section_content),
            "line_start": sections[current_section].get("line_start"),
            "line_end": len(lines) - 1
        }
    
    return sections

def extract_contact_info(text):
    """
    Extract contact information from resume text
    
    Args:
        text: Resume text
    
    Returns:
        dict: Extracted contact information
    """
    contact_info = {
        "email": None,
        "phone": None,
        "linkedin": None,
        "github": None,
        "website": None
    }
    
    # Email pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_matches = re.findall(email_pattern, text)
    if email_matches:
        contact_info["email"] = email_matches[0]
    
    # Phone pattern
    phone_pattern = r'(?:\+\d{1,3}[- ]?)?\(?(?:\d{3})?\)?[- ]?\d{3}[- ]?\d{4}'
    phone_matches = re.findall(phone_pattern, text)
    if phone_matches:
        contact_info["phone"] = phone_matches[0]
    
    # LinkedIn pattern
    linkedin_pattern = r'(?:linkedin\.com/in/|linkedin:)([A-Za-z0-9_-]+)'
    linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
    if linkedin_matches:
        contact_info["linkedin"] = f"linkedin.com/in/{linkedin_matches[0]}"
    
    # GitHub pattern
    github_pattern = r'(?:github\.com/|github:)([A-Za-z0-9_-]+)'
    github_matches = re.findall(github_pattern, text, re.IGNORECASE)
    if github_matches:
        contact_info["github"] = f"github.com/{github_matches[0]}"
    
    # Website pattern
    website_pattern = r'(?:https?://)?(?:www\.)?([A-Za-z0-9][-A-Za-z0-9]{0,62}(?:\.[A-Za-z0-9][-A-Za-z0-9]{0,62})+)'
    website_matches = re.findall(website_pattern, text)
    # Filter out common domains that are likely not personal websites
    filtered_matches = [m for m in website_matches if not any(common in m.lower() for common in ['linkedin', 'github', 'google', 'facebook', 'twitter'])]
    if filtered_matches:
        contact_info["website"] = filtered_matches[0]
    
    return contact_info

def generate_enhanced_resume_pdf(resume_data):
    """
    Generate an enhanced resume PDF from structured resume data
    
    Args:
        resume_data (dict): Structured resume data with sections
    
    Returns:
        bytes: PDF file as bytes
    """
    # Create an in-memory buffer for the PDF
    buffer = io.BytesIO()
    
    # Create the PDF document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.5*inch,
        rightMargin=0.5*inch,
        topMargin=0.5*inch,
        bottomMargin=0.5*inch
    )
    
    # Styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='Name',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=0.1*inch
    ))
    styles.add(ParagraphStyle(
        name='Contact',
        parent=styles['Normal'],
        fontSize=9,
        alignment=TA_CENTER,
        spaceAfter=0.2*inch
    ))
    styles.add(ParagraphStyle(
        name='SectionHeader',
        parent=styles['Heading2'],
        fontSize=12,
        spaceAfter=0.1*inch,
        spaceBefore=0.1*inch
    ))
    styles.add(ParagraphStyle(
        name='NormalBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=0.25*inch,
        firstLineIndent=-0.25*inch,
        spaceBefore=0.1*inch
    ))
    
    # Document elements
    elements = []
    
    # Personal information
    if 'Personal Info' in resume_data:
        personal_info = resume_data['Personal Info']
        name = personal_info.get('name', 'Your Name')
        elements.append(Paragraph(name, styles['Name']))
        
        # Contact information
        contact_parts = []
        for field in ['email', 'phone', 'location']:
            if field in personal_info and personal_info[field]:
                contact_parts.append(personal_info[field])
        
        contact_line = " | ".join(contact_parts)
        elements.append(Paragraph(contact_line, styles['Contact']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Summary/Objective
    if 'Objective/Resume Summary' in resume_data and resume_data['Objective/Resume Summary']:
        elements.append(Paragraph('SUMMARY', styles['SectionHeader']))
        elements.append(Paragraph(resume_data['Objective/Resume Summary'], styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Experience
    if 'Work Experience' in resume_data and resume_data['Work Experience']:
        elements.append(Paragraph('EXPERIENCE', styles['SectionHeader']))
        
        for job in resume_data['Work Experience']:
            job_title = job.get('title', 'Position')
            company = job.get('company', 'Company')
            date = job.get('duration', 'Date')
            
            job_header = f"<b>{job_title}</b>, {company} | {date}"
            elements.append(Paragraph(job_header, styles['Normal']))
            
            if 'description' in job:
                description = job['description']
                if isinstance(description, list):
                    # Handle list of bullet points
                    bullet_list = []
                    for bullet in description:
                        bullet_list.append(ListItem(Paragraph(bullet, styles['Normal'])))
                    elements.append(ListFlowable(bullet_list, bulletType='bullet', leftIndent=0.25*inch))
                else:
                    # Handle string description
                    elements.append(Paragraph(description, styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
    
    # Education
    if 'Education' in resume_data and resume_data['Education']:
        elements.append(Paragraph('EDUCATION', styles['SectionHeader']))
        
        for edu in resume_data['Education']:
            degree = edu.get('degree', 'Degree')
            institution = edu.get('institution', 'Institution')
            date = edu.get('date', 'Date')
            
            edu_header = f"<b>{degree}</b>, {institution} | {date}"
            elements.append(Paragraph(edu_header, styles['Normal']))
            
            if 'details' in edu and edu['details']:
                elements.append(Paragraph(edu['details'], styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
    
    # Skills
    if 'Skills' in resume_data and resume_data['Skills']:
        elements.append(Paragraph('SKILLS', styles['SectionHeader']))
        
        skills = resume_data['Skills']
        if isinstance(skills, dict):
            # Skills are categorized
            for category, skill_list in skills.items():
                elements.append(Paragraph(f"<b>{category}:</b> {', '.join(skill_list)}", styles['Normal']))
                elements.append(Spacer(1, 0.05*inch))
        elif isinstance(skills, list):
            # Skills are a simple list
            elements.append(Paragraph(', '.join(skills), styles['Normal']))
        else:
            # Skills is a string
            elements.append(Paragraph(skills, styles['Normal']))
    
    # Projects
    if 'Projects' in resume_data and resume_data['Projects']:
        elements.append(Paragraph('PROJECTS', styles['SectionHeader']))
        
        for project in resume_data['Projects']:
            name = project.get('name', 'Project Name')
            date = project.get('date', '')
            
            if date:
                project_header = f"<b>{name}</b> | {date}"
            else:
                project_header = f"<b>{name}</b>"
                
            elements.append(Paragraph(project_header, styles['Normal']))
            
            if 'description' in project:
                description = project['description']
                if isinstance(description, list):
                    # Handle list of bullet points
                    bullet_list = []
                    for bullet in description:
                        bullet_list.append(ListItem(Paragraph(bullet, styles['Normal'])))
                    elements.append(ListFlowable(bullet_list, bulletType='bullet', leftIndent=0.25*inch))
                else:
                    # Handle string description
                    elements.append(Paragraph(description, styles['Normal']))
            
            elements.append(Spacer(1, 0.1*inch))
    
    # Certifications
    if 'Certifications' in resume_data and resume_data['Certifications']:
        elements.append(Paragraph('CERTIFICATIONS', styles['SectionHeader']))
        
        for cert in resume_data['Certifications']:
            if isinstance(cert, dict):
                name = cert.get('name', 'Certification')
                issuer = cert.get('issuer', '')
                date = cert.get('date', '')
                
                cert_line = f"<b>{name}</b>"
                if issuer:
                    cert_line += f", {issuer}"
                if date:
                    cert_line += f" | {date}"
                
                elements.append(Paragraph(cert_line, styles['Normal']))
            else:
                elements.append(Paragraph(cert, styles['Normal']))
        
        elements.append(Spacer(1, 0.1*inch))
    
    # Build the PDF
    doc.build(elements)
    
    # Get the PDF data
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

def generate_docx_from_pdf(pdf_data):
    """
    Generate a DOCX version of the resume from PDF data
    
    Args:
        pdf_data (bytes): PDF file as bytes
    
    Returns:
        bytes: DOCX file as bytes
    """
    # Create temporary files
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as pdf_temp:
        pdf_temp.write(pdf_data)
        pdf_path = pdf_temp.name
    
    docx_path = pdf_path.replace('.pdf', '.docx')
    
    try:
        # In a real implementation, you would use a library like pdf2docx
        # For this example, we'll just create a minimal DOCX
        doc = docx.Document()
        doc.add_heading('Resume', 0)
        doc.add_paragraph('This is a placeholder DOCX generated from a PDF resume.')
        
        # Save the document
        doc.save(docx_path)
        
        # Read the DOCX file
        with open(docx_path, 'rb') as docx_file:
            docx_data = docx_file.read()
        
        return docx_data
    
    except Exception as e:
        print(f"Error converting PDF to DOCX: {str(e)}")
        return None
    
    finally:
        # Clean up temporary files
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path) 